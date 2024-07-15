#requirement. to enable Windows long path support
#https://pip.pypa.io/warnings/enable-long-paths
#python -m pip install --upgrade pip setuptools wheel
#python -m pip install --upgrade emoji certifi cffi charset-normalizer ffmpeg-python filelock future huggingface-hub idna lxml more-itertools numpy nvidia-cublas-cu11 nvidia-cuda-nvrtc-cu11 nvidia-cuda-runtime-cu11 nvidia-cudnn-cu11 openai-whisper packaging Pillow progressbar pycparser pydub python-docx PyYAML regex requests srt tokenizers torch tqdm transformers typing_extensions urllib3 vosk websockets 

import os
import re
import json
import shutil
import emoji # New import, required to remove non standard chars 

from PIL import Image
from glob import glob

from docx import Document
from docx.shared import Inches
import progressbar
from progressbar import Percentage as Perc

from vosk import Model, KaldiRecognizer, SetLogLevel
from pydub import AudioSegment


SetLogLevel(-1)


class Percentage(Perc):
    def update(self, pbar):
        return '%.2f%%' % pbar.percentage()


def unpacked() -> "tuple(str)": #makes it work
    cwd = os.getcwd()
    archive_files = glob(cwd + '/' + '*.zip')
    whatsapp_archive = None
    
    for file in archive_files:
        if 'whatsapp' in file.lower():
            whatsapp_archive = file

    if not whatsapp_archive:
        print('File archive WhatsApp not found')
        return None, None
    
    extract_dir = cwd + "/temp_data"
    shutil.unpack_archive(whatsapp_archive, extract_dir)
    print('WhatsApp archive unpack well done!')

    return extract_dir, whatsapp_archive.split('/')[-1].split('.')[0]


def image_to_jpg(image_path: str) -> str:
    jpg_image_path = image_path + '.jpg'
    Image.open(image_path).convert('RGB').save(jpg_image_path)
    return jpg_image_path


def convert_to_docx(file_name_txt: str, whatsapp_extract_dir: str, result_file_name: str) -> None:
    img_data = {'jpg', 'webp','png', 'jfif', 'exif', 'gif', 'tiff', 'bmp'}
    document = Document()

    with open(file_name_txt) as f:
        for i in f:
            if not i=='\n' #doesnt like empty lines, so remove them.
                suffix = i.split()[-1][:-1].split('.')[-1]
                if suffix in img_data:
                    document.add_paragraph(i)
                    filename = whatsapp_extract_dir + '/' + i.split()[-1][:-1]
                    convert_image = image_to_jpg(filename)
                    document.add_picture(convert_image, width=Inches(6))
                else:
                    document.add_paragraph(i)
        
        document.save(os.getcwd() + '/' + result_file_name + '.docx')

    os.remove(file_name_txt)
    shutil.rmtree(whatsapp_extract_dir)
    print('Temp files removed!')
    print('\33[92m' + '!!!DONE!!!')


def voice_recognition(filename: str) -> str:

    FRAME_RATE = 16000
    CHANNELS = 1

    # model = Model(model_path="vosk-model-ru-0.22")
    model = Model(model_name="vosk-model-small-ru-0.22")

    rec = KaldiRecognizer(model, FRAME_RATE)
    rec.SetWords(True)

    sound = AudioSegment.from_file(filename, codec="opus")

    sound = sound.set_channels(CHANNELS)
    sound = sound.set_frame_rate(FRAME_RATE)

    rec.AcceptWaveform(sound.raw_data)
    result = rec.Result()

    text = json.loads(result)["text"]
    return text


def run(whatsapp_extract_dir: str) -> str:
    try:
        with open(whatsapp_extract_dir+'/_chat.txt', 'r', errors="ignore") as file: #makes it work
            file_size = len(file.read())

    except FileNotFoundError as e:
        print('Error! File not found in "_chat.txt" !')
        return

     with open(replace_string(whatsapp_extract_dir+'/_chat.txt'), 'r', encoding='utf8', errors="ignore") as f, open('result.txt', 'w') as f2: #non standard characters cause issues, I brute force removed them
        step = 0
        bar = progressbar.ProgressBar(
            maxval=file_size,
            widgets=[
                progressbar.Bar(
                    marker='█',
                    left='[',
                    right=']',
                    # marker='\33[92m'+'█',
                    # left='\33[34m'+'[',
                    # right='\33[34m'+']',
                    fill='_'),
                    ' ',
                    progressbar.AnimatedMarker(markers='◐◓◑◒'),
                    ' ',
                    Percentage(),
                    ' ',
                    progressbar.ETA(),
                    ' ',
                    progressbar.FileTransferSpeed()
            ]
        )
        bar.start()
        for i in f:
            step += len(i)
            if re.findall('.opus' , i.lower()):
                filename = whatsapp_extract_dir + '/' + i.split()[-1][:-1]
                text_rec = voice_recognition(filename)
                i = i + '\t' + '-> ' + text_rec + '\n'
            f2.write(replace_string(i)) #non standard characters cause issues, I brute force removed them
            bar.update(step)
        bar.finish()
        print('All files transcribe well done!')
        
        return 'result.txt'
         
def replace_string(text): #I'm a brute, add any characters from forign charsets here
    text = emoji.replace_emoji(text,replace='')
    text = text.replace(u"\u202a",'')
    text = text.replace(u"\u202c",'')
    text = text.replace(u"\u202f",'')
    text = text.replace(u"\u200e",'')
    text = text.replace(u"\u015e",'S')
    text = text.replace(u"\u015f",'s')
    text = text.replace(u"\u0308",'O')
    text = text.replace(u"\u0327",'O')
    text = text.replace(u"\u0130",'i')
    text = text.replace(u"\u0131",'i')
    return text


def main() -> None:
    whatsapp_extract_dir, result_file_name = unpacked()
    if whatsapp_extract_dir:
        result_txt = run(whatsapp_extract_dir)
        convert_to_docx(result_txt, whatsapp_extract_dir, result_file_name)
    

if __name__ == '__main__':
    main()
    
